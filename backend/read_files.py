import os
import asyncio
from concurrent.futures import ThreadPoolExecutor
import fitz  # PyMuPDF for PDF handling
import docx
from PIL import Image
import io
import tempfile
import pytesseract
import re
import logging
import fitz
from docx import Document
from typing import Dict, List, Tuple
import time


# Configure pytesseract path (update based on your system)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Example for Windows

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class ReadFiles:
    """
    Class to extract text from PDF and DOCX files, including text from embedded images,
    with parallel processing for pages and images.
    """
    def __init__(self, max_workers=None):
        self.executor = ThreadPoolExecutor(max_workers=max_workers or max(os.cpu_count() * 2, 4))
        self.loop = asyncio.get_event_loop()
        self.semaphore = asyncio.Semaphore(16)

    async def get_text_from_image(self, image_data):
        """
        Extract text from an image using pytesseract asynchronously.

        Parameters:
        ---------
        image_data: Bytes or PIL Image object.

        Return:
        ------
        str: Extracted text.
        """
        try:
            if isinstance(image_data, bytes):
                image = Image.open(io.BytesIO(image_data)).convert('RGB')
            elif isinstance(image_data, Image.Image):
                image = image_data.convert('RGB')
            else:
                raise ValueError("Unsupported image data type")

            text = await self.loop.run_in_executor(
                self.executor, lambda: pytesseract.image_to_string(image, config='--psm 6')
            )
            return re.sub(r'\s+', ' ', text).strip()

        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""

    async def get_text_docx(self, file_path):
        """
        Extract text from a DOCX file, including paragraphs, tables, and embedded images.

        Parameters:
        ---------
        file_path: Path to the DOCX file.

        Return:
        ------
        str: Extracted text.
        """
        try:
            logger.debug(f"Processing DOCX: {file_path}")
            doc = await self.loop.run_in_executor(self.executor, lambda: docx.Document(file_path))
            full_text = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

            # Extract images from DOCX
            with tempfile.TemporaryDirectory() as temp_dir:
                import zipfile
                with zipfile.ZipFile(file_path) as docx_zip:
                    image_tasks = []
                    for file_info in docx_zip.infolist():
                        if file_info.filename.startswith('word/media/'):
                            image_data = docx_zip.read(file_info)
                            image_tasks.append(self.get_text_from_image(image_data))

                    if image_tasks:
                        image_texts = await asyncio.gather(*image_tasks, return_exceptions=True)
                        for text in image_texts:
                            if isinstance(text, str) and text.strip():
                                full_text.append(text)
                                logger.debug(f"Extracted text from DOCX image: {text[:100]}...")

            return ' '.join(full_text).strip()

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""

    async def get_text_pdf_page(self, page, page_num):
        """
        Extract text from a single PDF page (native text only).

        Parameters:
        ---------
        page: PyMuPDF page object.
        page_num: Page number (for logging).

        Return:
        ------
        str: Extracted text or empty string if none.
        """
        async with self.semaphore:
            try:
                start_time = time.time()
                text = await self.loop.run_in_executor(
                    self.executor, lambda: page.get_text("text", flags=fitz.TEXTFLAGS_TEXT).replace("\n", " ").replace(" -", "-")
                )
                text = re.sub(r'\s+', ' ', text).strip()
                logger.debug(f"Page {page_num + 1} processed in {time.time() - start_time:.2f}s, text length: {len(text)}")
                return text
            except Exception as e:
                logger.error(f"Error processing PDF page {page_num + 1}: {e}")
                return ""

    async def extract_pdf_images(self, pdf_document):
        """
        Extract text from embedded images in a PDF using PyMuPDF.

        Parameters:
        ---------
        pdf_document: PyMuPDF document object.

        Return:
        ------
        list: List of extracted text from images.
        """
        try:
            image_tasks = []
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        image_bytes = pdf_document.extract_image(xref)["image"]
                        task = self.get_text_from_image(image_bytes)
                        image_tasks.append(task)
                        logger.debug(f"Found image {img_index + 1} on page {page_num + 1}")
                    except Exception as e:
                        logger.error(f"Error extracting image {img_index + 1} from page {page_num + 1}: {e}")

            if image_tasks:
                image_texts = await asyncio.gather(*image_tasks, return_exceptions=True)
                return [text for text in image_texts if isinstance(text, str) and text.strip()]
            return []

        except Exception as e:
            logger.error(f"Error extracting images from PDF: {e}")
            return []

    async def get_text_pdf(self, file_content: io.BytesIO):
        """
        Extract text from a PDF file in-memory, processing pages in parallel, ignoring embedded images.

        Parameters:
        ---------
        file_content: BytesIO object containing PDF data.

        Return:
        ------
        str: Extracted text or empty string if none.
        """
        try:
            start_time = time.time()
            logger.debug("Processing PDF in-memory")
            with fitz.open(stream=file_content, filetype="pdf") as pdf:
                total_pages = len(pdf)
                logger.debug(f"Total pages: {total_pages}")
                if total_pages == 0:
                    logger.warning("PDF has no pages")
                    return ""

                # Batch pages to reduce task overhead
                batch_size = 50
                all_text = []
                for start in range(0, total_pages, batch_size):
                    end = min(start + batch_size, total_pages)
                    page_tasks = [
                        self.get_text_pdf_page(pdf[page_num], page_num)
                        for page_num in range(start, end)
                    ]
                    page_texts = await asyncio.gather(*page_tasks, return_exceptions=True)
                    all_text.extend([t for t in page_texts if isinstance(t, str) and t.strip()])

                combined_text = ' '.join(all_text).strip()
                logger.info(f"PDF text extraction took {time.time() - start_time:.2f}s, text length: {len(combined_text)}")
                return combined_text if combined_text else ""
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return ""

    async def process_file(self, filename: str, file_content: io.BytesIO) -> Tuple[str, str]:
        """
        Process a single file (PDF or DOCX) in-memory.

        Parameters:
        ---------
        filename: Name of the file.
        file_content: BytesIO object containing file data.

        Return:
        ------
        tuple: (filename, extracted_text)
        """
        try:
            file_ext = filename.split(".")[-1].lower()
            logger.debug(f"Processing file: {filename} ({file_ext})")
            if file_ext == "pdf":
                text = await self.get_text_pdf(file_content)
            elif file_ext in ["doc", "docx"]:
                doc = await self.loop.run_in_executor(self.executor, lambda: Document(file_content))
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            else:
                logger.error(f"Unsupported file format: {file_ext}")
                return filename, ""
            logger.debug(f"Extracted text from {filename}: {len(text)} characters")
            return filename, text
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            return filename, ""

    async def file_reader(self, files: List[Tuple[str, io.BytesIO]]) -> Dict[str, str]:
        """
        Extract text from multiple files (PDF, DOC, DOCX) in parallel, in-memory.

        Parameters:
        ---------
        files: List of tuples (filename, BytesIO object).

        Returns:
        -------
        Dict[str, str]: Dictionary mapping filenames to extracted text.
        """
        try:
            results = {}
            tasks = [self.process_file(filename, content) for filename, content in files]
            results_list = await asyncio.gather(*tasks, return_exceptions=True)
            for filename, text in results_list:
                results[filename] = text  # Include empty text
            return results
        except Exception as e:
            logger.error(f"Error in file_reader: {str(e)}")
            raise

    def __del__(self):
        """Clean up ThreadPoolExecutor."""
        self.executor.shutdown(wait=True)